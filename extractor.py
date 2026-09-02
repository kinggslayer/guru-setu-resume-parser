import json
import logging
import re
import time
import random
import pdfplumber

# pdfminer logs a warning for every malformed font it meets. A folder of a
# few thousand CVs produces tens of thousands of identical lines, which
# buries real errors in the Streamlit log and slows the run down. The
# warnings are cosmetic — text still extracts — so they are silenced.
for _noisy in ("pdfminer", "pdfminer.pdffont", "pdfminer.pdfinterp",
               "pdfminer.pdfpage", "pdfminer.converter", "pdfplumber"):
    logging.getLogger(_noisy).setLevel(logging.ERROR)
from docx import Document
from openai import OpenAI
from dotenv import load_dotenv
import os
import io
import base64

from usage import tracker

from portal_vocab import (
    GRADE_LEVELS,
    SUBJECTS,
    LANGUAGES,
    JOB_TYPES,
    AVAILABILITY,
    SKILLS,
    format_allowed,
    snap_grade_levels,
    snap_subjects,
    snap_languages,
    snap_qualification,
    snap_job_type,
    snap_availability,
    snap_gender,
)

load_dotenv()

client = OpenAI(
    api_key=os.getenv("OPENAI_API_KEY")
)

OPENAI_MODEL = "gpt-4.1-mini"

# Below this many characters a document can't plausibly be a resume — it's a
# scan, a cover image, or a failed read. Tuned low so a genuinely terse
# one-page CV still gets through.
MIN_USABLE_CHARS = 150

# How much of the resume is sent to the model.
#
# Name, contact, qualifications and current role are near the top of every
# CV; the tail is job-description prose we extract nothing from. Trimming
# 8000 -> 5000 chars cuts about 750 input tokens per resume with no
# measured loss, and read_pdf stops reading at this point too.
RESUME_CHAR_LIMIT = 5000

def sanitize_age(value):
    """
    The LLM sometimes returns a birthdate, a range, or other garbage in
    the 'age' field instead of a plain integer. Only accept something
    that's actually a plausible human age; everything else becomes None
    rather than polluting the data with junk like "January 02, 1987".
    Always returns a real Python int (never a float/str), or None.
    """

    if value is None:
        return None

    text = str(value).strip()

    try:
        age_int = int(float(text))

        if 15 <= age_int <= 80:
            return age_int

    except (ValueError, TypeError):
        pass

    # "34 years", "Age: 34", "34 yrs" — a plain int() fails on all of these
    # and the age was silently lost. Take the first plausible number, but
    # only when the text is short: a sentence could contain any number.
    if len(text) <= 30:

        for match in re.findall(r"\d{1,3}", text):

            candidate = int(match)

            if 15 <= candidate <= 80:
                return candidate

    return None


def sanitize_experience(value):
    """Coerce experience_years to a real int (not float/str), or None."""

    if value is None:
        return None

    try:
        return int(float(str(value).strip()))

    except (ValueError, TypeError):
        return None


def sanitize_salary(value):
    """
    Salary must be a plain number. Strips currency symbols, commas and any
    trailing unit text the model left behind. Returns None for anything that
    isn't a positive number, so a junk value never reaches the portal.
    """

    if value is None or value == "":
        return None

    cleaned = re.sub(r"[^\d.]", "", str(value))

    if not cleaned or cleaned.count(".") > 1:
        return None

    try:
        amount = float(cleaned)

    except (ValueError, TypeError):
        return None

    if amount <= 0:
        return None

    # Whole rupee figures shouldn't carry a meaningless ".0" into the sheet.
    return int(amount) if amount == int(amount) else round(amount, 2)


def sanitize_salary_unit(value, amount=None):
    """
    Only 'pm' or 'lpa'. When the model omits the unit we infer it from the
    magnitude: a teaching salary of 45 is lakhs per annum, 45000 is monthly.
    """

    text = str(value or "").strip().lower()

    if "lpa" in text or "lakh" in text or "annum" in text or "year" in text:
        return "lpa"

    if "pm" in text or "month" in text:
        return "pm"

    if amount is None:
        return None

    return "lpa" if float(amount) < 1000 else "pm"


def reconcile_salary(amount, unit):
    """
    Make the amount agree with its unit, and return (amount, unit).

    The portal reads 'lpa' as a figure IN LAKHS, so 450000 with unit 'lpa'
    means 450000 lakh — about 45 billion rupees. The model is asked to
    convert, but it returns the raw annual figure often enough that this
    has to be enforced here rather than trusted.

    The reverse is also caught: 4.5 marked 'pm' is not a monthly salary in
    rupees, it is lakhs per annum.
    """

    if amount is None:
        return None, None

    value = float(amount)

    if unit == "lpa" and value >= 1000:
        # An annual rupee figure, not lakhs. 450000 -> 4.5
        value = round(value / 100000, 2)

    elif unit == "pm" and value < 1000:
        # Too small to be a monthly wage; it is lakhs per annum.
        return (int(value) if value == int(value) else round(value, 2)), "lpa"

    if value <= 0:
        return None, None

    return (int(value) if value == int(value) else round(value, 2)), unit


def sanitize_college_type(value):
    """
    Only 'IIT', 'NIT', or None are valid — guard against model drift.

    Matched on word boundaries, because a substring test classified "IIIT
    Hyderabad" as an IIT and "MANIT" or "SNIT" as an NIT. Those are
    different institutions, and the portal filters on this field.
    """

    if not value:
        return None

    value_upper = str(value).strip().upper()

    if re.search(r"\bIIT\b", value_upper):
        return "IIT"

    if re.search(r"\bNIT\b", value_upper):
        return "NIT"

    return None


def sanitize_phone(value):
    """
    Normalize every phone number to the SAME format: +91XXXXXXXXXX.
    Handles numbers that already have +91, 91, spaces, dashes, or nothing.
    Returns None if it doesn't look like a valid 10-digit Indian mobile
    number after cleanup, rather than keeping garbage.
    """

    if not value:
        return None

    text = str(value)

    digits = re.sub(r"\D", "", text)

    normalised = _as_indian_mobile(digits)

    if normalised:
        return normalised

    # Resumes often list two numbers: "9812345678 / 9876543210" or
    # "+91-98123 45678, 98765 43210". Stripping all non-digits then glues
    # them into one 20-digit string that matches nothing, and the candidate
    # ends up with NO phone at all. Fall back to scanning for the first
    # number that stands on its own.
    for candidate in re.findall(r"\d[\d\s\-().]{7,}\d", text):

        normalised = _as_indian_mobile(re.sub(r"\D", "", candidate))

        if normalised:
            return normalised

    return None


def _as_indian_mobile(digits):
    """digits -> +91XXXXXXXXXX, or None if it isn't a valid mobile."""

    if len(digits) == 12 and digits.startswith("91"):
        digits = digits[2:]

    if len(digits) == 11 and digits.startswith("0"):
        digits = digits[1:]

    if len(digits) == 10 and digits[0] in "6789":
        return f"+91{digits}"

    return None


def flatten_value(item):
    """
    The LLM is inconsistent about whether education_history/college/
    previous_institutions entries are plain strings or nested objects
    with arbitrary keys. This forces everything into a single readable
    string so every record in the cache/Excel has the same shape.
    """

    if isinstance(item, str):
        text = item.strip()

        # The literal words too: earlier versions wrote str(None) into
        # these fields, so "None" and "nan" appear in real stored data.
        return text if text.lower() not in ("none", "nan", "null") else None

    if isinstance(item, dict):
        parts = [
            f"{k}: {v}"
            for k, v in item.items()
            if v not in (None, "", [])
        ]
        return "; ".join(parts) if parts else None

    if isinstance(item, list):
        flattened = [flatten_value(x) for x in item]
        flattened = [x for x in flattened if x]
        return " / ".join(flattened) if flattened else None

    # None must come back as None, not the string "None".
    #
    # str(None) is "None", which is truthy, so it survived every downstream
    # emptiness check and was written into education_history, college and
    # previous_institutions as if it were real data. The same mistake
    # elsewhere merged unrelated candidates and crashed a whole run.
    if item is None:
        return None

    text = str(item).strip()

    return text if text.lower() not in ("none", "nan", "null") else None


def flatten_list_field(value):
    """Apply flatten_value to every item in a list field, drop empties."""

    if not isinstance(value, list):
        return []

    result = [flatten_value(item) for item in value]
    return [item for item in result if item]


# A resume is a handful of pages. Reading beyond this is wasted work on a
# mis-filed document, and the memory it costs is what kills a large run.
MAX_PDF_PAGES = 12


def read_pdf(file, char_limit=RESUME_CHAR_LIMIT, max_pages=MAX_PDF_PAGES):
    text = ""

    with pdfplumber.open(file) as pdf:

        for index, page in enumerate(pdf.pages):

            if index >= max_pages:
                break

            page_text = page.extract_text()

            if page_text:
                text += page_text + "\n"

            # pdfplumber keeps each page's parsed layout in memory for the
            # lifetime of the document. Across thousands of CVs on a 1 GB
            # container that accumulation is fatal, so release it per page.
            try:
                page.flush_cache()
                page.get_textmap.cache_clear()

            except Exception:
                pass

            if len(text) >= char_limit:
                break

    return text


# --------------------------------------------------------------------------
#  OCR for scanned resumes
#
#  A scanned or photographed CV has no text layer, so pdfplumber returns
#  nothing. Rather than skipping those candidates, the pages are rendered to
#  images and transcribed by the same vision-capable model already in use.
#
#  This deliberately produces TEXT and hands it back to the normal pipeline,
#  so the regex fallbacks, qualification parsing and content-hash dedup all
#  behave exactly as they do for a text PDF.
# --------------------------------------------------------------------------

# Images are billed by area, so each OCR page costs several times a page of
# text. Page 1 carries the name and contact details; page 2 catches an
# overflowing education table. Beyond that is rarely worth the money.
OCR_MAX_PAGES = 2

# 1.6x lands near 115 DPI. Still comfortably legible for a phone number,
# and the image is about 35% smaller in area than at 2x — which is close to
# a 35% saving on the most expensive call the app makes.
OCR_RENDER_SCALE = 1.6


def render_pdf_pages(file_path, max_pages=OCR_MAX_PAGES, scale=OCR_RENDER_SCALE):
    """Render the first pages of a PDF to base64 PNGs."""

    import pypdfium2

    images = []

    pdf = pypdfium2.PdfDocument(file_path)

    try:

        for index in range(min(len(pdf), max_pages)):

            bitmap = pdf[index].render(scale=scale)
            image = bitmap.to_pil()

            buffer = io.BytesIO()
            image.save(buffer, format="PNG")

            images.append(
                base64.b64encode(buffer.getvalue()).decode("utf-8")
            )

    finally:
        pdf.close()

    return images


def ocr_pdf(file_path):
    """
    Transcribe a scanned PDF to plain text. Returns "" on any failure, so
    the caller just sees an unreadable document rather than an exception.
    """

    # MOCK_MODE bypassed the extraction call but not this one, so a test
    # run still made real (and comparatively expensive) vision requests.
    if os.getenv("MOCK_MODE", "false").lower() == "true":
        return "MOCK OCR TEXT. " * 40

    try:
        images = render_pdf_pages(file_path)

    except Exception as e:
        print(f"[OCR] Could not render {file_path}: {e}", flush=True)
        return ""

    if not images:
        return ""

    content = [
        {
            "type": "text",
            "text": (
                "Transcribe this resume exactly as written, preserving the "
                "reading order. Include every email address, phone number, "
                "qualification, school name and date. Output only the "
                "transcribed text, with no commentary and no markdown."
            )
        }
    ]

    for image in images:
        content.append({
            "type": "image_url",
            "image_url": {"url": f"data:image/png;base64,{image}"}
        })

    try:

        response = client.chat.completions.create(
            model=OPENAI_MODEL,
            messages=[{"role": "user", "content": content}],
            temperature=0,
            max_tokens=2000
        )

        tracker.record(response)

        return response.choices[0].message.content or ""

    except Exception as e:
        print(f"[OCR] Transcription failed: {type(e).__name__}: {e}", flush=True)
        return ""


def _iter_block_text(container):
    """
    Yield text from paragraphs AND tables, recursing into nested tables.

    python-docx's .paragraphs only covers top-level body paragraphs — text
    inside a table cell is invisible to it. Plenty of teacher resumes put
    personal details and education in a table, so reading paragraphs alone
    silently loses the email, phone and qualifications.
    """

    for para in container.paragraphs:

        if para.text.strip():
            yield para.text

    for table in container.tables:

        for row in table.rows:

            cells = []

            for cell in row.cells:

                # A cell is itself a block container, so nested tables and
                # multi-paragraph cells come through this recursion.
                cell_text = " ".join(_iter_block_text(cell))

                if cell_text.strip():
                    cells.append(cell_text.strip())

            if cells:
                # Tab-separated keeps "Email <tab> x@y.com" on one line, which
                # reads naturally for both the regexes and the LLM.
                yield "\t".join(cells)


def read_docx(file):

    doc = Document(file)

    parts = list(_iter_block_text(doc))

    # Contact details are often parked in the header or footer.
    for section in doc.sections:

        for area in (section.header, section.footer):

            try:
                parts.extend(_iter_block_text(area))

            except Exception:
                continue

    return "\n".join(parts)


def extract_email(text):

    match = re.search(
        r'[\w\.-]+@[\w\.-]+\.\w+',
        text
    )

    return match.group(0) if match else None


def extract_phone(text):

    match = re.search(
        r'(\+91[- ]?)?[6-9]\d{9}',
        text
    )

    return match.group(0) if match else None


def extract_experience(text):

    matches = []

    matches.extend(
        re.findall(
            r'(\d+)\+?\s*years',
            text,
            re.IGNORECASE
        )
    )

    matches.extend(
        re.findall(
            r'(\d+)\+?\s*yrs',
            text,
            re.IGNORECASE
        )
    )

    matches.extend(
        re.findall(
            r'total\s*yrs?.*?(\d+)',
            text,
            re.IGNORECASE
        )
    )

    if matches:

        try:
            return max(
                int(x)
                for x in matches
            )

        except Exception:
            pass

    return None


# Each entry: canonical name -> the spellings that mean it.
#
# Matching is by regex with word boundaries, NOT substring containment.
# A plain "mba" in text.lower() matches inside "Bo(mba)y", so every
# candidate who studied at IIT Bombay was recorded as holding an MBA.
_QUALIFICATION_PATTERNS = [
    ("PhD", [r"ph\.?\s?d\b", r"doctor of philosophy", r"\bdoctorate\b"]),
    ("UGC NET", [r"ugc[\s-]*net\b"]),
    ("CSIR-NET", [r"csir[\s-]*net\b"]),
    ("NET", [r"\bnet\b"]),
    ("JRF", [r"\bjrf\b"]),
    ("GATE", [r"\bgate\b"]),

    ("MBA", [r"\bm\.?\s?b\.?\s?a\b", r"master of business administration"]),
    ("PGDM", [r"\bpgdm\b"]),
    ("M.Tech", [r"\bm\.?\s?tech\b", r"master of technology",
                r"master'?s? of engineering", r"\bm\.?\s?e\.?\s*\("]),
    ("M.Ed", [r"\bm\.?\s?ed\b", r"master of education"]),
    ("M.Sc", [r"\bm\.?\s?sc\b", r"master of science",
              r"master'?s?\s+(?:of|in)\s+\w+"]),
    ("M.A", [r"\bm\.?\s?a\b", r"master of arts"]),
    ("M.Com", [r"\bm\.?\s?com\b", r"master of commerce"]),
    ("MCA", [r"\bmca\b", r"master of computer applications?"]),
    ("MBM", [r"\bmbm\b"]),

    ("B.Tech", [r"\bb\.?\s?tech\b", r"bachelor of technology",
                r"\bb\.?\s?e\.?\s*\(", r"bachelor of engineering"]),
    ("B.Ed", [r"\bb\.?\s?ed\b", r"bachelor of education"]),
    ("B.Sc", [r"\bb\.?\s?sc\b", r"bachelor of science"]),
    ("B.A", [r"\bb\.?\s?a\b", r"bachelor of arts"]),
    ("B.Com", [r"\bb\.?\s?com\b", r"bachelor of commerce"]),
    ("BCA", [r"\bbca\b", r"bachelor of computer applications?"]),
    ("BBA", [r"\bbba\b"]),

    ("D.El.Ed", [r"\bd\.?\s?el\.?\s?ed\b", r"\bdeled\b"]),
    ("D.T.Ed", [r"\bd\.?\s?t\.?\s?ed\b"]),
    ("DSM", [r"\bdsm\b"]),

    ("CTET", [r"\bctet\b"]),
    ("TET", [r"\b(?<!c)tet\b", r"\bstate\s+tet\b"]),
]

# Where one qualification's text contains another's, listing both is noise:
# "CTET qualified" should not also record a bare "TET".
_QUALIFICATION_SUPERSEDES = {
    "CTET": ["TET"],
    "UGC NET": ["NET"],
    "CSIR-NET": ["NET"],
}


def extract_qualifications(text):
    """
    Qualifications named in the resume, most specific first.

    Regex with word boundaries rather than substring search — see the note
    on _QUALIFICATION_PATTERNS for why that matters.
    """

    lower_text = str(text or "").lower()

    found = []

    for canonical, patterns in _QUALIFICATION_PATTERNS:

        for pattern in patterns:

            if re.search(pattern, lower_text):
                found.append(canonical)
                break

    # Drop the broader match when a more specific one is present.
    for specific, broader in _QUALIFICATION_SUPERSEDES.items():

        if specific in found:

            for term in broader:

                if term in found:
                    found.remove(term)

    return found


_QUALIFICATION_RANK = {
    "PhD": 100,
    "Ph.D": 100,
    "Ph. D": 100,
    "Doctor of Philosophy": 100,
    "UGC NET": 95,
    "CSIR-NET": 95,
    "NET": 95,
    "JRF": 95,

    "GATE": 90,

    "MBA": 85,
    "PGDM": 85,
    "M.Tech": 85,
    "MCA": 85,

    "MBM": 80,
    "M.Ed": 80,
    "M.Sc": 80,
    "M.A": 80,
    "M.Com": 80,

    "B.Tech": 70,
    "B.Ed": 70,
    "B.Sc": 70,
    "B.A": 70,
    "B.Com": 70,
    "BCA": 70,
    "BBA": 70,

    "D.El.Ed": 60,
    "D.T.Ed": 60,
    "DSM": 50,

    "CTET": 40,
    "TET": 40
}


def split_qualifications(qualification_list):

    if not qualification_list:
        return None, []

    ranked = sorted(
        qualification_list,
        key=lambda q: _QUALIFICATION_RANK.get(q, 0),
        reverse=True
    )

    highest = ranked[0]
    extras = [q for q in ranked[1:] if q != highest]

    return highest, extras


def extract_resume_data(text):

    if os.getenv("MOCK_MODE", "false").lower() == "true":
        return {
            "full_name": "Test User",
            "gender": "Male",
            "age": 25,
            "city": "Delhi",

            "state": "Delhi",

            "subjects": ["Mathematics"],
            "grade_levels": ["Senior Secondary (11-12)"],
            "languages": ["English"],

            "college_type": None,
            "college": ["University of Delhi"],
            "education_history": ["B.Sc Mathematics, University of Delhi"],

            "current_institution": "ABC School",
            "current_designation": "PGT Mathematics",
            "previous_institutions": ["XYZ School"],

            "experience_years": 3,
            "preferred_job_type": "Full-time",
            "availability": "Within 1 month",

            "current_salary": 40000,
            "current_salary_unit": "pm",
            "expected_salary": 6.0,
            "expected_salary_unit": "lpa",

            "email": "test@example.com",
            "phone": "+919999999999",

            "qualification": "M.Sc",
            "extra_qualifications": [],

            "extraction_failed": False
        }

    short_text = text[:RESUME_CHAR_LIMIT]
    lower_text = text.lower()

    # A scanned or image-only resume extracts to (almost) nothing. Without
    # this guard the model is handed an empty prompt and invents a plausible
    # candidate, which is worse than no record at all — and it costs a call.
    if len(text.strip()) < MIN_USABLE_CHARS:
        return {
            "extraction_failed": True,
            "needs_ocr": True,
            "error": (
                f"Only {len(text.strip())} characters of text could be read — "
                "this looks like a scanned or image-only document. It needs "
                "OCR before it can be parsed."
            )
        }

    if (
        "offer letter" in lower_text
        or "appointment as" in lower_text
    ):
        return {
            "extraction_failed": True,
            "error": "Document appears to be an offer letter, not a resume"
        }

    grade_levels_allowed = format_allowed(GRADE_LEVELS)
    subjects_allowed = format_allowed(SUBJECTS)
    languages_allowed = format_allowed(LANGUAGES)
    job_types_allowed = format_allowed(JOB_TYPES)
    availability_allowed = format_allowed(AVAILABILITY)
    skills_allowed = format_allowed(SKILLS)

    instructions = f"""
You are an expert teacher resume parser. Extract structured data from the resume text below.

Return ONLY valid JSON, matching this EXACT schema. No explanations, no markdown, no extra keys.

{{
    "full_name": null,
    "gender": null,
    "age": null,
    "city": null,
    "state": null,
    "subjects": [],
    "grade_levels": [],
    "languages": [],
    "skills": [],
    "college_type": null,
    "college": [],
    "education_history": [],
    "current_institution": null,
    "current_designation": null,
    "previous_institutions": [],
    "experience_years": null,
    "preferred_job_type": null,
    "availability": null,
    "current_salary": null,
    "current_salary_unit": null,
    "expected_salary": null,
    "expected_salary_unit": null
}}

CRITICAL RULES — follow these exactly:

=== CONTROLLED VOCABULARY (most important rules) ===

- "grade_levels" must contain ONLY values copied EXACTLY from this list:
  {grade_levels_allowed}
  Translate whatever the resume says onto these values. Examples:
    "11th & 12th", "HSC", "PGT", "+2"      -> "Senior Secondary (11-12)"
    "9th and 10th", "TGT", "high school"   -> "Secondary (9-10)"
    "IIT-JEE", "engineering entrance"      -> "JEE Mains" AND "JEE Advanced"
    "IIT-JEE Mains only"                   -> "JEE Mains"
    "medical entrance", "AIPMT"            -> "NEET"
    "MHT-CET", "KCET"                      -> "State CET"
    "CBSE board classes"                   -> "Board Exams"
    "dropper batch", "repeaters"           -> "Dropper / Repeater"
    "B.Sc students", "degree college"      -> "College / UG"
    "nursery", "KG"                        -> "Pre-Primary"
  One resume can map to several values — list all that apply. Do NOT invent
  values outside the list. If the resume never says which classes the person
  teaches, return an empty list.

- "subjects" must contain ONLY values copied EXACTLY from this list:
  {subjects_allowed}
  Map specialist topics onto their parent subject ("Calculus" -> "Mathematics",
  "Optics" -> "Physics", "Genetics" -> "Biology"). Put the original specialist
  wording in "skills" instead so nothing is lost. Never put soft skills or
  generic words like "Teaching" here. Empty list if not clearly stated.

- "languages" must contain ONLY values copied EXACTLY from this list:
  {languages_allowed}
  Ignore programming languages here — those belong in "skills".

- "preferred_job_type" must be EXACTLY one of: {job_types_allowed}, or null.
  Only set it if the resume actually states a preference. Do not assume
  "Full-time" just because the person has a job.

- "availability" must be EXACTLY one of: {availability_allowed}, or null.
  Derive it from a stated notice period or joining date:
    "can join immediately"      -> "Immediate"
    "15 days notice"            -> "Within 15 days"
    "1 month notice"            -> "Within 1 month"
    "2 months notice"           -> "Within 2 months"
    "3 months notice", "serving notice" -> "Notice period"
  Return null if the resume says nothing about availability.

- "skills" are technical/professional skills. Prefer these standard values
  where they apply: {skills_allowed}
  You may also add resume-specific ones (e.g. "MS-CIT", "Tally Prime",
  "C++", "LaTeX", "Python") and the specialist topics mentioned above.

=== SALARY ===

- "current_salary" and "expected_salary" must be plain NUMBERS with no
  currency symbol, commas or units (450000, not "4.5 LPA" or "Rs. 4,50,000").
- "current_salary_unit" / "expected_salary_unit" must be EXACTLY "pm"
  (a per-MONTH figure) or "lpa" (LAKHS per annum), matching how the resume
  states it:
    "Rs 35,000 per month"  -> current_salary 35000,  unit "pm"
    "CTC 4.5 LPA"          -> current_salary 4.5,    unit "lpa"
    "Annual salary 540000" -> current_salary 5.4,    unit "lpa"
  For "lpa", give the figure IN LAKHS (4.5, not 450000).
  Leave all four null if the resume does not state salary — never guess.

=== LOCATION ===

- "state" must be the Indian state or union territory the city belongs to,
  spelled in full ("Maharashtra", "Uttar Pradesh", "Delhi", "Karnataka").
  You may derive it from the city when the city is unambiguous (Pune ->
  Maharashtra, Indore -> Madhya Pradesh). Return null if the city is unknown
  or genuinely ambiguous. Never put a district, taluka or locality here.


- "age" must be a plain integer (e.g. 28), representing the person's current age in years.
  NEVER put a date of birth, a year, or any text in "age". If the resume doesn't state an
  actual age (a number of years), leave "age" as null. If DOB or birth year is present,
calculate current age in years. Return a plain integer age.
Examples:
02-09-2002 -> 23
23/02/1974 -> 52
If no DOB or age information exists, return null.
- "gender" — if not explicitly stated in the resume, INFER it from the person's first name
  using common Indian naming conventions (e.g. "Priya" -> Female, "Rajesh" -> Male). Return
  "Male" or "Female". Only leave this null if the name is genuinely ambiguous or unavailable
  (e.g. only initials given, like "A. Kumar").
- "college_type" must be EXACTLY one of: "IIT", "NIT", or null.
  Set it to "IIT" only if the person studied (any degree) at any Indian Institute of
  Technology. Set it to "NIT" only if they studied at any National Institute of Technology.
  If neither applies, leave it null — do NOT put any other institution name, degree type,
  or descriptive text in this field.
- college must contain ONLY the institutions
associated with the candidate's highest
educational qualifications.
Maximum 3 institutions.
Do NOT include:
- SSC schools
- HSC schools
- coaching centres
- certification institutes
- every institution from education history
Examples:
GOOD:
["IIT Delhi"]
GOOD:
["Yashwantrao Chavan University",
 "North Maharashtra University"]
BAD:
["School",
 "College",
 "SSC Board",
 "Coaching Centre",
 "University",
 "Training Institute"] — plain strings like
  "Indian Institute of Technology, Bombay" or "University of Delhi". Do NOT put degree
  names, years, or percentages in this field (e.g. "B.Sc (2021)" is WRONG — that belongs
  in education_history, not here).
- city must be a single city/town name only.

- Do NOT guess, infer, correct, or construct city names.
- current_designation should be the person's most recent
  job title/designation.

- Common examples:
  Faculty
  Biology Faculty
  Mathematics Faculty
  Physics Faculty
  Chemistry Faculty
  Teacher
  Lecturer
  Professor
  Assistant Professor
  Principal
  Coordinator
  Trainer

- If the resume mentions "Faculty",
  "Teaching Assistant",
  "Professor",
  "Lecturer",
  etc., do not leave current_designation blank.

- Return null only when no designation can be identified.
- Do NOT return locality names, colony names,
  streets, areas, villages, districts,
  addresses, landmarks, or combinations of locations.

- If multiple locations are present,
  return only the actual city/town name.

Examples:

GOOD:
"Mumbai"
"Dhule"
"Buldana"
"Ichalkaranji"
"Delhi"

BAD:
"Shivaji Nagar Korochi"
"Sector 62 Noida"
"Near Bus Stand"
"Taluka Hatkanangale"
"Korochi, Kolhapur"
"Village XYZ"

- If a clear city/town name cannot be determined,
  return null.
- education_history should include only graduation,
  post-graduation, diploma, doctorate,
  professional qualifications, certifications,
  and higher education.
- previous_institutions must contain only places
  where the candidate WORKED.
- If an institution appears in education_history,
  it should NOT appear in previous_institutions
  unless the resume explicitly states that the
  candidate worked there.
- Do NOT include:
  colleges
  universities
  schools attended as a student
  training institutes
  certification institutes

- Include only organizations where the candidate
  was employed, taught, trained students,
  or held a professional role.

Examples:

GOOD:
Delhi Public School
Aakash Institute
FIITJEE
Allen Career Institute

BAD:
IIT Delhi
University of Delhi
North Maharashtra University
Jai Hind College
Canossa Convent High School
- Exclude SSC, HSC, Class 10,
  Class 12, Secondary Education,
  Higher Secondary Education entries.
- "education_history" must be a list of plain strings, one per degree/qualification,
  each combining degree + institution + year into ONE readable string, e.g.
  "M.Sc. Chemistry, XYZ University, 2020, 75%". Do not return nested objects.
- "experience_years" must be a plain integer — the person's TOTAL years of relevant
  teaching/professional work experience. Calculate this by looking at the work history
  section: if date ranges are given (e.g. "2019 - Present", "2015 - 2018"), compute the
  total span of relevant experience. If the resume explicitly states a number of years
  of experience, use that instead. If neither can be determined confidently, leave null
  — do NOT guess a number.
- "current_institution" must be the name of the school/institution the person CURRENTLY
  works at (their most recent employer), not a college.
- "current_designation" must be the person's CURRENT or MOST RECENT job title only
  (e.g. "Senior Physics Faculty", "PGT Chemistry Teacher") — not a list, not a summary
  of their whole career, just the one current/latest title.
- "previous_institutions" must be a list of plain strings — the names of EMPLOYERS
  (schools/institutions the person has WORKED at, not studied at) prior to their current
  one, e.g. ["Delhi Public School, Noida", "Ryan International School"]. Do NOT include
  the current institution in this list, and do NOT include colleges/universities here —
  those belong in "college". Leave empty if the resume shows no prior employer or only
  one job overall.
- Do not guess or hallucinate any value except gender as instructed above.
  If something else isn't clearly stated, use null (or an empty list for list fields).
- languages must be a list of plain strings.
- Return valid JSON only — nothing before or after it.
"""

    # The resume is the ONLY variable part, kept in its own message so the
    # instructions above stay a byte-identical prefix across every call.
    resume_block = f"Resume:\n\n{short_text}"

    parsed = {}
    extraction_failed = False

    max_attempts = 4
    base_delay = 2
    last_error = None

    for attempt in range(max_attempts):

        try:

            response = client.chat.completions.create(
                model=OPENAI_MODEL,
                messages=[
                    # System holds ONLY the static instructions, so the
                    # cacheable prefix is identical on every call and
                    # OpenAI's automatic prompt caching applies from the
                    # second resume onwards. Mixing the resume in here
                    # would break the prefix and forfeit the discount.
                    {
                        "role": "system",
                        "content": instructions
                    },
                    {
                        "role": "user",
                        "content": resume_block
                    }
                ],
                temperature=0,
                # Measured completions run ~450 tokens. This is a runaway
                # guard, not a target; output is billed at 4x input.
                max_tokens=1200,
                response_format={"type": "json_object"}
            )

            tracker.record(response)

            content = (
                response.choices[0]
                .message.content
                .replace("```json", "")
                .replace("```", "")
                .strip()
            )

            json_match = re.search(r"\{.*\}", content, re.DOTALL)

            if json_match:
                content = json_match.group(0)

            parsed = json.loads(content)

            if isinstance(parsed, list):

                if len(parsed) > 0:
                    parsed = parsed[0]
                else:
                    parsed = {}

            break

        except Exception as e:

            print(
                f"[OpenAI] Error on attempt: {type(e).__name__}: {e}",
                flush=True
            )

            err_str = str(e)
            last_error = f"{type(e).__name__}: {err_str}"

            lower_err = err_str.lower()

            # A 429 usually means "slow down", which is worth retrying. But
            # an exhausted balance, a revoked key or a missing model also
            # surface as errors that will NEVER succeed on a retry — backing
            # off four times just makes the app look frozen for minutes.
            is_permanent = any(
                marker in lower_err
                for marker in (
                    "insufficient_quota",
                    "credit_balance_exhausted",
                    "billing_hard_limit_reached",
                    "invalid_api_key",
                    "incorrect api key",
                    "account_deactivated",
                    "model_not_found",
                    "does not exist or you do not have access",
                )
            )

            is_rate_limit = (
                not is_permanent
                and ("429" in err_str or "rate_limit" in lower_err)
            )

            is_parse_error = isinstance(e, json.JSONDecodeError)

            if is_permanent:
                extraction_failed = True
                break

            if (is_rate_limit or is_parse_error) and attempt < max_attempts - 1:

                retry_after = None

                match = re.search(
                    r"try again in ([\d.]+)s",
                    err_str,
                    re.IGNORECASE
                )

                if match:
                    retry_after = float(match.group(1))

                delay = retry_after if retry_after else (base_delay * (2 ** attempt))
                delay += random.uniform(0, 1)

                time.sleep(delay)
                continue

            else:
                extraction_failed = True
                break

    else:
        extraction_failed = True

    qualification_list = extract_qualifications(text)

    highest_qualification, extra_qualification_list = split_qualifications(qualification_list)

    parsed["email"] = extract_email(text)
    parsed["phone"] = extract_phone(text)

    llm_experience = sanitize_experience(parsed.get("experience_years"))

    parsed["experience_years"] = llm_experience if llm_experience is not None else extract_experience(text)

    parsed["qualification"] = highest_qualification
    parsed["extra_qualifications"] = extra_qualification_list

    parsed.setdefault("full_name", None)
    parsed.setdefault("gender", None)
    parsed.setdefault("age", None)
    parsed.setdefault("city", None)
    parsed.setdefault("state", None)
    parsed.setdefault("preferred_job_type", None)
    parsed.setdefault("availability", None)

    parsed.setdefault("subjects", [])
    parsed.setdefault("grade_levels", [])
    parsed.setdefault("languages", [])
    parsed.setdefault("skills", [])

    parsed.setdefault("college_type", None)

    parsed.setdefault("college", [])
    parsed.setdefault("education_history", [])

    parsed.setdefault("current_institution", None)
    parsed.setdefault("current_designation", None)
    parsed.setdefault("previous_institutions", [])
    parsed["age"] = sanitize_age(parsed.get("age"))
    parsed["experience_years"] = sanitize_experience(parsed.get("experience_years"))
    parsed["college_type"] = sanitize_college_type(parsed.get("college_type"))

    parsed["phone"] = sanitize_phone(parsed.get("phone"))

    parsed["college"] = flatten_list_field(parsed.get("college"))
    parsed["education_history"] = flatten_list_field(parsed.get("education_history"))
    parsed["subjects"] = flatten_list_field(parsed.get("subjects"))
    parsed["grade_levels"] = flatten_list_field(parsed.get("grade_levels"))
    parsed["languages"] = flatten_list_field(parsed.get("languages"))
    parsed["skills"] = flatten_list_field(parsed.get("skills"))
    parsed["previous_institutions"] = flatten_list_field(parsed.get("previous_institutions"))

    # ---------------------------------------------------------------
    #  Snap everything onto the portal's controlled vocabulary.
    #  The prompt already asks for canonical values; this catches drift.
    #  Anything unrecognised is preserved rather than dropped:
    #    - unmatched grade levels  -> tags   (portal keeps them searchable)
    #    - unmatched subjects      -> skills (they're specialist topics)
    #    - unmatched languages     -> tags
    # ---------------------------------------------------------------

    grade_levels, grade_leftover = snap_grade_levels(parsed.get("grade_levels"))
    subjects, subject_leftover = snap_subjects(parsed.get("subjects"))
    languages, language_leftover = snap_languages(parsed.get("languages"))

    parsed["grade_levels"] = grade_levels
    parsed["subjects"] = subjects
    parsed["languages"] = languages

    parsed["skills"] = list(
        dict.fromkeys(parsed.get("skills", []) + subject_leftover)
    )

    parsed["tags"] = list(
        dict.fromkeys(
            parsed.get("tags", []) + grade_leftover + language_leftover
        )
    )

    parsed["gender"] = snap_gender(parsed.get("gender"))
    parsed["qualification"] = snap_qualification(parsed.get("qualification"))
    parsed["extra_qualifications"] = [
        snap_qualification(q)
        for q in parsed.get("extra_qualifications", [])
        if snap_qualification(q)
    ]

    parsed["preferred_job_type"] = snap_job_type(parsed.get("preferred_job_type"))
    parsed["availability"] = snap_availability(parsed.get("availability"))

    for field in ("current_salary", "expected_salary"):

        amount = sanitize_salary(parsed.get(field))
        unit = sanitize_salary_unit(parsed.get(f"{field}_unit"), amount)

        amount, unit = reconcile_salary(amount, unit)

        parsed[field] = amount
        parsed[f"{field}_unit"] = unit

    parsed["extraction_failed"] = extraction_failed

    # Why it failed, so the app can show it instead of just "failed".
    if extraction_failed and last_error:
        parsed["extraction_error"] = last_error

    return parsed